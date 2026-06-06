"""
generate_assignment.py
Generates Assignment 1 — NLP Semester Project Proposal (Milestone 1)
as a professionally formatted Word .docx document.

Course : CSC-355 Natural Language Processing
Student: [Your Name] | [Your Roll Number]
Run    : python generate_assignment.py
Output : Assignment1_NLP_Milestone1.docx  (in the same directory)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
#  HELPER UTILITIES
# ─────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line_spacing=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_horizontal_rule(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_heading(doc, number, title, level=1):
    """Bold, blue numbered section heading."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=4)
    run = p.add_run(f"Section {number}: {title}")
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
    run.font.name = 'Calibri'
    return p


def add_sub_heading(doc, title):
    """Smaller bold sub-heading (dark gray)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x24, 0x44, 0x6A)
    run.font.name = 'Calibri'
    return p


def add_body(doc, text, bold_parts=None):
    """Body paragraph with justified alignment."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Simple: just one run for the whole text
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def add_body_runs(doc, runs_spec):
    """
    Add a body paragraph with mixed bold/normal runs.
    runs_spec: list of (text, bold) tuples.
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in runs_spec:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text, indent_level=0):
    """Bullet list item."""
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.15)
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * indent_level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def add_bullet_runs(doc, runs_spec, indent_level=0):
    """Bullet item with mixed bold/normal runs."""
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.15)
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * indent_level)
    for text, bold in runs_spec:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_code_block(doc, code_text):
    """Monospace code block with light-gray shading."""
    # Add shaded paragraph for each line
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        # Gray shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)

        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Tiny spacer after block
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=4)


def add_table(doc, headers, rows, col_widths=None):
    """Simple bordered table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        # Blue header shading
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
        # Alternating row shading
        if r_idx % 2 == 1:
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EBF3FB')
                tcPr.append(shd)

    if col_widths:
        for r in table.rows:
            for idx, cell in enumerate(r.cells):
                cell.width = Inches(col_widths[idx])
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type='page')


# ─────────────────────────────────────────────
#  DOCUMENT CREATION
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page Margins ──────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # ─────────────────────────────────────────
    #  HEADER BLOCK
    # ─────────────────────────────────────────
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run("NAMAL UNIVERSITY MIANWALI")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run("Department of Computer Science")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    add_horizontal_rule(doc)

    # Course metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Course Code & Title:", "CSC-355 — Natural Language Processing"),
        ("Instructor:", "Dr. Muzamil Ahmed"),
        ("Session / Semester:", "2022–2026 | 8th Semester"),
        ("Total Marks:", "15"),
        ("Submission Date:", "19-03-2026"),
    ]
    for i, (label, value) in enumerate(meta_data):
        cells = meta_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10.5)
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()  # spacer

    # Assignment title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=4)
    run = p.add_run("Assignment 1 — NLP Semester Project Proposal")
    run.bold = True; run.font.size = Pt(14); run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6)
    run = p.add_run("Milestone 1: Project Conceptualization and Preprocessing Design")
    run.bold = True; run.font.size = Pt(12); run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # ─────────────────────────────────────────
    #  SECTION 1 — PROJECT TITLE
    # ─────────────────────────────────────────
    add_section_heading(doc, 1, "Proposed Project Title")
    add_horizontal_rule(doc)

    # Title in a box-style paragraph
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DEEAF1')
    pPr.append(shd)
    run = p.add_run(
        '"Sentiment Analysis of Code-Switched Urdu-English Social Media Text '
        'Using Multilingual Transformer Models"'
    )
    run.bold = True; run.italic = True
    run.font.size = Pt(12); run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    add_body(doc,
        "The proposed title has been formulated with deliberate precision to communicate the full scope, "
        "linguistic challenge, task type, and computational methodology of this research initiative. Each "
        "component of the title performs a specific semantic function: the phrase 'Sentiment Analysis' "
        "identifies the NLP task (opinion mining and polarity classification); 'Code-Switched Urdu-English' "
        "specifies the unique and underexplored linguistic phenomenon being addressed; 'Social Media Text' "
        "situates the domain within a high-volume, noisy, real-world data source; and 'Multilingual "
        "Transformer Models' signals the advanced computational methodology — namely, multilingual pretrained "
        "language models such as mBERT and XLM-RoBERTa — that will be employed at the culmination of the "
        "project pipeline."
    )

    add_body(doc,
        "The title is intentionally specific and solution-oriented, avoiding generic labels such as 'Text "
        "Analysis System' or 'NLP Classification Project.' It reflects a semester-long technical trajectory: "
        "beginning at Milestone 1 with text preprocessing and baseline model design, progressing through "
        "word embeddings (FastText, Word2Vec) and sequence models (BiLSTM) in subsequent milestones, and "
        "culminating in fine-tuned multilingual Transformer models evaluated with rigorous NLP metrics. "
        "This trajectory directly aligns with CLO-3 of the course — 'Design NLP solutions using neural "
        "networks, Transformers, and pretrained models' — ensuring that the project title is not merely "
        "a label, but a forward-looking commitment to the full course curriculum."
    )

    # ─────────────────────────────────────────
    #  SECTION 2 — PROBLEM STATEMENT
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 2, "Problem Statement")
    add_horizontal_rule(doc)

    add_body(doc,
        "Pakistani social media platforms — including Twitter/X, Facebook, and Instagram — are characterised "
        "by a pervasive linguistic phenomenon known as code-switching, wherein users seamlessly intersperse "
        "Urdu words and phrases within English sentences, and vice versa, within a single post or comment. "
        "A representative example is: 'Yeh movie bohat achi thi, definitely watch karo! 🔥', which blends "
        "Urdu lexicons ('yeh', 'bohat', 'achi', 'karo') with English vocabulary within a single coherent "
        "utterance. This produces text that is syntactically ambiguous, phonetically transliterated as "
        "Roman Urdu, and linguistically inconsistent at the structural level. Existing NLP sentiment "
        "analysis tools — almost universally trained on monolingual English corpora — produce near-random "
        "classification results on such mixed-language input, rendering automated opinion mining dangerously "
        "unreliable for Pakistani digital markets, brand monitoring systems, and public discourse analysis platforms."
    )

    add_body(doc,
        "The difficulty of this problem is compounded by several interacting factors that collectively "
        "render standard NLP approaches wholly inadequate. Code-switched text has no fixed grammar, no "
        "standardised spelling conventions for Roman Urdu (the same Urdu word may appear in dozens of "
        "orthographic variants, e.g., 'acha', 'acha', 'accha', 'acha'), and exhibits extremely high "
        "lexical variability. Sentiment polarity can shift mid-sentence between the two language segments, "
        "introducing cross-lingual semantic ambiguity. Standard tokenisers — including those in NLTK and "
        "spaCy — incorrectly segment transliterated Urdu text; standard English stopword lists miss "
        "high-frequency Urdu function words entirely; and word embeddings trained exclusively on English "
        "corpora assign zero-vectors or arbitrary representations to Urdu tokens, fundamentally degrading "
        "the quality of any downstream model trained on such features."
    )

    add_body(doc,
        "The necessity of an automated NLP solution cannot be overstated. Pakistan generates millions of "
        "social media posts daily, and manual annotation of even a modest sample at scale would be "
        "computationally and economically infeasible. Only automated, learned NLP pipelines — capable of "
        "generalising across unseen, noisy, code-switched inputs — can operate at the production scale "
        "demanded by commercial and research applications. The computational challenge is therefore "
        "substantial: any viable system must simultaneously handle two typologically distinct language "
        "families (Indo-Aryan Urdu and Germanic English), resolve transliteration ambiguity, process "
        "emoji semantics as sentiment-bearing signals, and learn cross-lingual sentiment representations "
        "that no rule-based system could feasibly encode — thereby mandating the use of deep neural "
        "language models as the primary computational paradigm."
    )

    # ─────────────────────────────────────────
    #  SECTION 3 — MOTIVATION
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 3, "Motivation")
    add_horizontal_rule(doc)

    add_body(doc,
        "Pakistan counts over 70 million active social media users, a number that continues to grow "
        "rapidly across urban and rural populations alike. This scale creates an urgent, commercially "
        "significant demand for automated sentiment analysis tools capable of processing the unique "
        "linguistic reality of Pakistani social media discourse. Businesses require robust brand monitoring "
        "systems that can aggregate and classify customer sentiment expressed in code-switched text; "
        "political analysts require real-time public opinion tracking on policy decisions articulated in "
        "mixed Urdu-English; academic researchers require labelled sentiment corpora to advance the "
        "field of South Asian computational linguistics; and technology startups building Urdu-language "
        "AI products require foundational NLP pipelines on which to build downstream applications. All "
        "of these use cases are currently blocked by the absence of high-quality, code-switching-aware "
        "sentiment analysis tools."
    )

    add_body(doc,
        "From a research perspective, Urdu-English code-switched sentiment analysis represents a "
        "significantly under-explored area in the global NLP literature. The dominant NLP benchmarks "
        "— including GLUE, SuperGLUE, and SemEval — are overwhelmingly English-centric, and most "
        "multilingual benchmarks address high-resource language pairs. Building robust, open-source "
        "NLP tools for Urdu-English code-switched text would constitute a meaningful contribution to "
        "NLP equity, and would establish methodological templates directly applicable to analogous "
        "code-switching contexts worldwide — including Hindi-English, Arabic-English, Swahili-English, "
        "and Mandarin-English mixed-language communities. This project therefore carries implications "
        "that extend far beyond the immediate Pakistani context. Furthermore, this initiative directly "
        "instantiates the course's emphasis on real-world NLP system development, as articulated in "
        "Unit 6 of the course outline, which addresses the design, evaluation, and deployment of "
        "full-stack NLP systems in authentic, domain-specific settings."
    )

    # ─────────────────────────────────────────
    #  SECTION 4 — EXPECTED OUTCOMES
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 4, "Expected Outcomes")
    add_horizontal_rule(doc)

    add_body(doc,
        "The following five measurable and verifiable outcomes are expected upon completion of all five "
        "project milestones. Each outcome maps directly to the corresponding milestone in the course "
        "schedule (Lectures 10, 16, 22, 27, and 30 respectively):"
    )

    outcomes = [
        (
            "Outcome 1 — Text Preprocessing Pipeline (Milestone 1, Due: Lecture 10): ",
            "A complete, modular text preprocessing pipeline — implemented in Python — capable of "
            "correctly handling Roman Urdu transliteration, Urdu Nastaliq Unicode script, emojis, "
            "Twitter @mentions, hashtags, URLs, bilingual stopwords, and code-switched tokens. "
            "The pipeline will be designed to serve as the data preparation layer for all subsequent "
            "model experiments across Milestones 2 through 5."
        ),
        (
            "Outcome 2 — Labelled Sentiment Dataset (Milestone 2, Due: Lecture 16): ",
            "A curated and labelled dataset of at least 10,000 code-switched Urdu-English social media "
            "posts with sentence-level sentiment annotations (Positive, Negative, Neutral), combining "
            "the publicly available SentiMix dataset (SemEval-2020 Task 9) with a self-collected corpus "
            "of Pakistani Twitter/X posts gathered via the academic API and manually annotated using "
            "an inter-annotator agreement protocol."
        ),
        (
            "Outcome 3 — Comparative Multi-Model Evaluation (Milestone 3, Due: Lecture 22): ",
            "A rigorous comparative evaluation of at least four models spanning the full modelling "
            "spectrum covered in the course: (i) Logistic Regression with TF-IDF features (classical "
            "baseline); (ii) BiLSTM with FastText multilingual embeddings (neural sequence model); "
            "(iii) multilingual BERT (mBERT) with fine-tuning; and (iv) XLM-RoBERTa with fine-tuning — "
            "evaluated using Macro F1-score, Precision, Recall, Accuracy, and Confusion Matrix."
        ),
        (
            "Outcome 4 — Fine-Tuned Transformer Model (Milestone 4, Due: Lecture 27): ",
            "A fine-tuned mBERT or XLM-RoBERTa model achieving a minimum target Macro F1-score of "
            "≥ 0.78 on the held-out test set, with documented hyperparameter tuning, learning rate "
            "scheduling, and class-imbalance mitigation strategies applied during the fine-tuning process."
        ),
        (
            "Outcome 5 — Final Evaluation and Deployment Report (Milestone 5, Due: Lecture 30): ",
            "A comprehensive written evaluation report documenting: model performance comparisons "
            "across all four architectures; an error analysis of mispredicted samples; a bias and "
            "fairness assessment; an estimation of computational resource requirements; and a feasibility "
            "analysis for deploying the fine-tuned model as a lightweight inference API or mobile "
            "on-device service."
        ),
    ]

    for label, body in outcomes:
        add_bullet_runs(doc, [(label, True), (body, False)])

    # ─────────────────────────────────────────
    #  SECTION 5 — DATASET IDENTIFICATION
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 5, "Dataset Identification")
    add_horizontal_rule(doc)

    add_sub_heading(doc, "5.1  Primary Dataset — SentiMix (SemEval-2020 Task 9)")
    add_body(doc,
        "The primary dataset for this project is SentiMix, the official corpus released for "
        "SemEval-2020 Task 9: Sentiment Analysis for Code-Mixed Social Media Text. This is a "
        "publicly available, peer-reviewed benchmark dataset comprising approximately 15,000 "
        "labelled tweets in English-Hindi code-switched text — the most linguistically proximate "
        "publicly available resource to Urdu-English code-switching, given that Hindi and Urdu "
        "share the same grammatical structure and substantial lexical overlap (both are derived from "
        "Khari Boli). Sentence-level sentiment annotations are provided, with three classes: "
        "Positive, Negative, and Neutral. The dataset is available at: "
        "https://competitions.codalab.org/competitions/20911. The use of a competitive benchmark "
        "dataset ensures reproducibility and enables direct comparison with published results in "
        "the SemEval proceedings."
    )

    add_sub_heading(doc, "5.2  Secondary Dataset — Self-Collected Pakistani Twitter/X Corpus")
    add_body(doc,
        "To supplement SentiMix with authentic Urdu-English code-switched content specific to "
        "Pakistani social media discourse, a secondary corpus of approximately 3,000–5,000 posts "
        "will be collected from Twitter/X using the Academic Research Product Track API, applying "
        "keyword filters for high-frequency Urdu-English mixed terms across domains including "
        "film reviews, political commentary, product feedback, and sports. Manual sentiment "
        "annotation will be performed using a standardised three-label scheme (Positive, Negative, "
        "Neutral), with inter-annotator agreement measured via Cohen's Kappa statistic. Language "
        "data includes both Roman Urdu script (transliterated) and Urdu Nastaliq Unicode script "
        "(U+0600–U+06FF). No token-level language identification labels are provided in the raw "
        "corpora; language identification per token will be performed as an automated preprocessing "
        "step using Unicode character range heuristics."
    )

    add_sub_heading(doc, "5.3  Dataset Quality and Generalization")
    add_body(doc,
        "Both dataset size and annotation quality directly govern the generalisation capacity of "
        "any supervised learning model trained on them, as discussed extensively in Yoav Goldberg's "
        "Neural Network Methods for Natural Language Processing. Insufficient or inconsistently "
        "labelled data produces models that memorise training artifacts rather than learning "
        "transferable sentiment representations — a risk that will be mitigated in this project "
        "through careful annotation protocols, dataset stratification, and the application of "
        "cross-lingual transfer learning via pretrained multilingual models that have been "
        "exposed to large-scale multilingual corpora during pretraining."
    )

    # ─────────────────────────────────────────
    #  SECTION 6 — PREPROCESSING STEPS
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 6, "Required NLP Text Preprocessing")
    add_horizontal_rule(doc)

    add_body(doc,
        "The preprocessing pipeline for code-switched Urdu-English social media text is substantially "
        "more complex than standard monolingual English pipelines. The following ten steps are required, "
        "each justified specifically with respect to the linguistic properties of the target data. Note "
        "that in Milestone 3, the custom tokenisation stage will be partially superseded by mBERT's "
        "built-in WordPiece subword tokeniser; however, the full pipeline below is essential for "
        "baseline models in Milestones 1 and 2."
    )

    preprocessing_steps = [
        (
            "Step 1 — Unicode Normalization",
            "Urdu text in digital form exists across multiple, visually identical but byte-level distinct "
            "Unicode representations — specifically NFC (Canonical Decomposition followed by Canonical "
            "Composition) and NFD (Canonical Decomposition). Without Unicode normalisation to a single "
            "canonical form, the identical Urdu word may be represented as two distinct byte sequences, "
            "causing it to be treated as two separate vocabulary items by any downstream tokeniser or "
            "embedding lookup. Unicode normalisation to NFC is therefore a non-negotiable prerequisite "
            "for any Urdu script processing pipeline."
        ),
        (
            "Step 2 — Lowercasing (English Tokens Only)",
            "Uppercase variation in English words ('Happy', 'HAPPY', 'happy') introduces artificial "
            "vocabulary expansion without semantic benefit. Lowercasing normalises these variants to "
            "a single type, reducing vocabulary size and improving embedding quality. Crucially, this "
            "operation affects only Latin-script (English and Roman Urdu) characters; Urdu Nastaliq "
            "script is inherently case-insensitive and is unaffected by this transformation."
        ),
        (
            "Step 3 — URL Removal",
            "Social media posts frequently contain hyperlinks (e.g., 'https://t.co/...') that carry "
            "zero intrinsic sentiment information but introduce highly variable token sequences that "
            "confuse both statistical and neural models. All URLs matching standard HTTP/HTTPS patterns "
            "are removed using regular expression filters before any tokenisation is performed."
        ),
        (
            "Step 4 — Mention and Hashtag Cleaning",
            "@username mentions are entirely uninformative for sentiment classification and must be "
            "removed. Hashtag symbols (#) are stripped, but the hashtag word itself is retained, "
            "as the content frequently carries sentiment-bearing meaning (e.g., '#Disappointed' → "
            "'disappointed'). This nuanced handling is specific to social media NLP pipelines and "
            "has no equivalent in standard document corpora."
        ),
        (
            "Step 5 — Emoji Handling",
            "Emojis are among the strongest sentiment signals in informal social media text. A naive "
            "approach of simply deleting emojis would discard this rich sentiment information entirely. "
            "Instead, emojis are converted to descriptive text labels using the Python emoji library's "
            "demojize() function (e.g., 😊 → 'smiling_face', 😢 → 'crying_face'), preserving their "
            "semantic content in a form that downstream models can process as standard tokens."
        ),
        (
            "Step 6 — HTML Entity and Noise Removal",
            "Web-scraped and API-collected social media text frequently contains HTML entities "
            "(&amp;, &lt;, &gt;), residual XML tags, and non-linguistic special characters. These "
            "must be removed using targeted regular expression patterns. Critically, the removal "
            "pattern must preserve the Urdu Unicode character range (U+0600–U+06FF) and "
            "sentiment-bearing punctuation marks (! and ?) to avoid inadvertent information loss."
        ),
        (
            "Step 7 — Custom Tokenisation",
            "Standard NLP tokenisers — including both NLTK's word_tokenize and spaCy's default "
            "tokeniser — are trained exclusively on English text and produce systematically incorrect "
            "segmentation for Urdu tokens and code-switched sequences. A hybrid tokenisation approach "
            "based on whitespace splitting supplemented by regex pattern matching is therefore employed "
            "for the baseline model stages. In later milestones (from Milestone 3 onwards), mBERT's "
            "multilingual WordPiece tokeniser — which has been trained on 104 languages including "
            "both Urdu and English — will provide a far superior subword-level segmentation that "
            "naturally handles out-of-vocabulary Urdu tokens."
        ),
        (
            "Step 8 — Bilingual Stopword Removal",
            "Standard English stopword lists (as provided by NLTK) contain no Urdu function words. "
            "In code-switched text, high-frequency Urdu function words such as 'ہے', 'ہیں', 'کا', "
            "'کی', 'اور' contribute no discriminative sentiment signal yet would otherwise consume "
            "valuable vocabulary capacity. A custom bilingual stopword list — combining NLTK's "
            "English stopwords with a curated Urdu stopword lexicon — is required to adequately "
            "reduce noise across both language components."
        ),
        (
            "Step 9 — Punctuation and Number Filtering",
            "Pure numeric tokens (stand-alone digit strings) carry no sentiment value in this context "
            "and are removed. Standard punctuation is filtered, but sentiment-bearing punctuation "
            "characters — most significantly the exclamation mark (!) and question mark (?) — are "
            "retained as potentially discriminative features, given their frequent association with "
            "strong emotional expression in informal social media writing."
        ),
        (
            "Step 10 — Language Identification per Token",
            "Identifying the language of each individual token — as either Urdu (UR) or English (EN) "
            "— is a preprocessing step that is entirely unique to code-switched NLP pipelines and has "
            "no direct analogue in monolingual processing. A Unicode character-range heuristic is "
            "applied: tokens containing characters within U+0600–U+06FF are classified as Urdu; tokens "
            "containing predominantly ASCII characters are classified as English. This per-token "
            "language tag enables language-conditional feature extraction and will serve as the basis "
            "for language-aware analysis in subsequent milestones, particularly when comparing how "
            "mBERT handles Urdu versus English sub-sequences."
        ),
    ]

    for title, explanation in preprocessing_steps:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=6, after=3, line_spacing=1.15)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.2)
        run_title = p.add_run(title + " — ")
        run_title.bold = True
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run_body = p.add_run(explanation)
        run_body.font.name = 'Calibri'
        run_body.font.size = Pt(11)

    # ─────────────────────────────────────────
    #  SECTION 7 — PYTHON PREPROCESSING CODE
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 7, "Python Preprocessing Code")
    add_horizontal_rule(doc)

    add_body(doc,
        "The following Python implementation provides a complete, modular, and fully commented "
        "preprocessing pipeline for code-switched Urdu-English social media text. All functions "
        "correspond directly to the preprocessing steps described in Section 6. Sample inputs and "
        "outputs from diverse text types are included in the demonstration block."
    )

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    run = p.add_run("7.1  Full Implementation")
    run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    import os as _os
    _code_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), 'pipeline_code.txt')
    with open(_code_path, 'r', encoding='utf-8') as _f:
        CODE = _f.read()

    add_code_block(doc, CODE)


    # Summary table of pipeline steps
    add_sub_heading(doc, "7.2  Pipeline Step Summary Table")
    add_body(doc,
        "The table below summarises each pipeline function, its purpose, and the specific reason "
        "it is necessary for code-switched Urdu-English NLP — as distinct from generic monolingual pipelines."
    )

    table_headers = ["Step", "Function", "General Purpose", "Code-Switch-Specific Necessity"]
    table_rows = [
        ("1", "normalize_unicode()", "Consistent character encoding", "Multiple Unicode representations of identical Urdu characters cause spurious vocabulary duplication"),
        ("2", "convert_emojis()", "Preserve sentiment from emojis", "Emojis are primary sentiment signals in informal Urdu-English social media text"),
        ("3", "remove_urls()", "Eliminate noise tokens", "URLs appear at high frequency in tweets; produce unique tokens with zero sentiment value"),
        ("4", "clean_mentions_hashtags()", "Remove/normalise social media markers", "Hashtag content often carries sentiment; stripping only the # symbol retains it"),
        ("5", "lowercase_english()", "Normalise English case", "Affects only Latin characters; Urdu Nastaliq is inherently case-insensitive"),
        ("6", "remove_noise()", "Strip HTML and special characters", "Must explicitly preserve Urdu Unicode range (U+0600–U+06FF) to avoid losing Urdu tokens"),
        ("7", "tokenize()", "Split text into tokens", "Standard tokenisers fail on Urdu; mBERT WordPiece replaces this in Milestone 3"),
        ("8", "remove_stopwords()", "Reduce vocabulary noise", "Requires bilingual stopword list — English-only lists miss all Urdu function words"),
        ("9", "filter_numbers_and_punct()", "Remove non-informative tokens", "Sentiment punct (! ?) retained as discriminative features for social media text"),
        ("10", "tag_language()", "Identify language per token", "Unique to code-switched NLP — enables language-conditional analysis in later milestones"),
    ]
    add_table(doc, table_headers, table_rows, col_widths=[0.4, 1.6, 1.6, 2.5])

    # ─────────────────────────────────────────
    #  SECTION 8 — CCP ATTRIBUTES
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 8, "Why This Is a Complex Computing Problem (CCP Analysis)")
    add_horizontal_rule(doc)

    add_body(doc,
        "According to the assignment guidelines, a qualifying complex computing problem must satisfy "
        "a defined set of attributes relating to data scale, computational non-triviality, ambiguity, "
        "algorithmic sophistication, evaluation rigour, real-world applicability, and multi-component "
        "system integration. The seven CCP attributes are addressed below with specific reference to "
        "this project:"
    )

    ccp_items = [
        (
            "1. Large-Scale Data Handling: ",
            "The primary corpus contains over 15,000 publicly sourced tweets supplemented by a "
            "self-collected dataset of 3,000–5,000 posts, requiring multi-stage preprocessing across "
            "heterogeneous data sources in two scripts and two language families."
        ),
        (
            "2. Non-Trivial Computation: ",
            "The project pipeline spans five progressively complex processing stages across the full "
            "semester — from rule-based preprocessing through classical machine learning, to neural "
            "sequence models (BiLSTM), to fine-tuning large pretrained Transformer models — none of "
            "which can be trivially implemented or substituted with off-the-shelf tools."
        ),
        (
            "3. Ambiguity Handling: ",
            "The system must simultaneously resolve code-switching boundaries, Roman Urdu "
            "orthographic ambiguity (multiple phonetic spellings per word), emoji polysemy, "
            "and cross-lingual sentiment polarity shifts — multiple interacting layers of "
            "linguistic ambiguity that demand learned, probabilistic representations rather "
            "than deterministic rule systems."
        ),
        (
            "4. Algorithmic Design: ",
            "The project requires the selection, design, and comparison of four distinct model "
            "architectures — TF-IDF with Logistic Regression, BiLSTM with FastText embeddings, "
            "fine-tuned mBERT, and fine-tuned XLM-RoBERTa — each embodying fundamentally different "
            "algorithmic assumptions about language representation and generalisation."
        ),
        (
            "5. Performance Evaluation: ",
            "Model performance will be assessed using a comprehensive suite of standard NLP "
            "metrics: Macro F1-score (primary metric for class-imbalanced data), Precision, "
            "Recall, Accuracy, and Confusion Matrix analysis — consistent with the evaluation "
            "framework presented in the Handbook of Natural Language Processing and covered in "
            "Unit 6 of the course outline."
        ),
        (
            "6. Real-World Constraints: ",
            "This project directly addresses an active, commercially relevant and unresolved "
            "problem in Pakistani social media analytics, operating under real data constraints "
            "including annotation cost, API rate limits, class imbalance, and the need for "
            "deployment-viable model sizes."
        ),
        (
            "7. Multiple Components Integration: ",
            "The final system integrates six distinct functional components: data collection and "
            "curation, custom multilingual preprocessing, classical feature engineering, neural "
            "model training with transfer learning, rigorous cross-architecture evaluation, and "
            "deployment feasibility analysis — constituting a full, end-to-end NLP system of "
            "substantial engineering and research complexity."
        ),
    ]

    for label, body in ccp_items:
        add_bullet_runs(doc, [(label, True), (body, False)])

    # ─────────────────────────────────────────
    #  EVALUATION RUBRIC TABLE
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 9, "Marking Rubric Reference")
    add_horizontal_rule(doc)

    add_body(doc,
        "The following rubric summarises the assessment components for this assignment as specified "
        "in the course outline, along with the sections of this report that address each component:"
    )

    rubric_headers = ["Component", "Marks", "Addressed In"]
    rubric_rows = [
        ("Project Title Quality — specific, technical, solution-oriented", "2", "Section 1"),
        ("Problem Statement — all four sub-points clearly addressed", "3", "Section 2"),
        ("Motivation — real impact, beneficiaries, research gap", "2", "Section 3"),
        ("Expected Outcomes — measurable, realistic, five outcomes listed", "2", "Section 4"),
        ("Dataset Identification — source, size, language, labels, availability", "2", "Section 5"),
        ("Preprocessing Design — ten steps each explained with justification", "2", "Section 6"),
        ("Code Quality — clean, commented, modular, with sample I/O and pipeline table", "2", "Section 7"),
        ("Technical Depth — CCP attributes addressed, course alignment demonstrated", "2", "Sections 1 & 8"),
        ("TOTAL", "15", "All Sections"),
    ]
    add_table(doc, rubric_headers, rubric_rows, col_widths=[3.8, 0.7, 1.6])

    # ─────────────────────────────────────────
    #  REFERENCES
    # ─────────────────────────────────────────
    doc.add_paragraph()
    add_section_heading(doc, 10, "References")
    add_horizontal_rule(doc)

    references = [
        "Lee, R. (2022). Natural Language Processing (1st ed.).",
        "Goldberg, Y. (2017). Neural Network Methods for Natural Language Processing (1st ed.). Morgan & Claypool Publishers.",
        "Kamath, U., Graham, K., & Chiu, W. (2024). Large Language Models: A Deep Dive. Springer.",
        "Dale, R., Moisl, H., & Somers, H. (Eds.). (2000). Handbook of Natural Language Processing (2nd ed.). Chapman & Hall/CRC.",
        "Patwa, P., Aguilar, G., Kar, S., Pandey, S., PYKL, S., Gambäck, B., Chakraborty, T., Solorio, T., & Das, A. (2020). "
        "SentiMix 2020: Sentiment Analysis for Code-Mixed Social Media Text. Proceedings of the Fifth Workshop on Computational "
        "Approaches to Linguistic Code-Switching, EMNLP 2020.",
        "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for "
        "Language Understanding. Proceedings of NAACL-HLT 2019.",
        "Conneau, A., et al. (2020). Unsupervised Cross-lingual Representation Learning at Scale. Proceedings of ACL 2020.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=4, line_spacing=1.15)
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)

    # ─────────────────────────────────────────
    #  SAVE
    # ─────────────────────────────────────────
    output_path = r"f:\Documents\Last Year\Semester 8\NPL\Project\Assignment1_NLP_Milestone1.docx"
    doc.save(output_path)
    print(f"\n[SUCCESS] Document saved to:\n  {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
