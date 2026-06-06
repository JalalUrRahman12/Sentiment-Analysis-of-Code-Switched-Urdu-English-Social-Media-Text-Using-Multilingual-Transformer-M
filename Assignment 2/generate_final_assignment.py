import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = docx.Document()

# --- STYLES ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# --- COVER PAGE ---
def add_center_para(text, size, bold=False, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

add_center_para("Namal University Mianwali", 20, bold=True, space_after=6)
add_center_para("Department of Computer Science", 16, bold=True, space_after=36)

add_center_para("Assignment 2: Related Work Review and Literature Analysis", 18, bold=True, space_after=24)
add_center_para("Course: CSC-355 Natural Language Processing", 14, space_after=6)
add_center_para("Instructor: Dr. Muzamil Ahmed", 14, space_after=24)

add_center_para("Project Title:", 14, bold=True, space_after=6)
add_center_para('"Sentiment Analysis of Code-Switched Urdu-English Social Media Text Using Multilingual Transformer Models"', 16, bold=True, space_after=36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted By:\n")
run.bold = True
run.font.size = Pt(14)
p.add_run("Jalal Ur Rahman (NUM-BSCS-2022-42)\nSultan Haider").font.size = Pt(14)
p.paragraph_format.space_after = Pt(36)

add_center_para("Date: 03-04-2026", 14, space_after=0)

doc.add_page_break()

# --- PARSE TEXT FILE ---
with open('assignment2_manus_output.txt', 'r', encoding='utf-8') as f:
    lines = [line.strip() for line in f.read().split('\n')]

table_data = []
in_table = False
table_lines = []

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        
for i, line in enumerate(lines):
    if not line:
        continue
        
    if line == "3. Comparative Literature Table":
        add_heading(line, 1)
        in_table = True
        continue
        
    if line == "4. Conclusion":
        in_table = False
        # Process accumulated table data
        num_cols = 7
        rows = [table_lines[i:i + num_cols] for i in range(0, len(table_lines), num_cols)]
        
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text
                # Bold header row
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            
        doc.add_paragraph() # Spacing
        add_heading(line, 1)
        continue

    if in_table:
        table_lines.append(line)
    else:
        # Check for headings
        if line == "Abstract" or line == "5. References" or line.startswith("1. Introduction") or line.startswith("2. Literature Review"):
            add_heading(line, 1)
        elif line.startswith("2.1.") or line.startswith("2.2.") or line.startswith("2.3."):
            add_heading(line, 2)
        elif line == "Sentiment Analysis of Code-Switched Urdu-English Social Media Text Using Multilingual Transformer Models: A Literature Review":
            # Document title
            add_center_para(line, 16, bold=True)
        else:
            # Normal paragraph
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()
add_heading("Appendix A: Plagiarism Report", 1)
doc.add_paragraph("[Please insert the screenshot or text of your Plagiarism Report here before creating the final PDF.]")

doc.add_page_break()
add_heading("Appendix B: AI Usage Report", 1)
doc.add_paragraph("[Please insert your AI Usage Report here before creating the final PDF.]")

output_path = 'Jalal_Sultan_Assignment2_Final.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
